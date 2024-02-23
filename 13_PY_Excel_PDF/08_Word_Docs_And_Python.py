# Word docoments and Python 

# Word docs have the .docx file extension, we'll 
# be using the third party module python-docx:

# pip install python-docx

# Compared to plain .txt files, .docx have a lot 
# more structure 

# Docx has three main data types:

# - Document Object: respresents the entire doc,
#   it contains Parahraph Objects, Paragraph objects 
#   contain Run objects

import docx

d = docx.Document('.\\demo.docx')

print(d.paragraphs[0].text)
# Document Title
print(d.paragraphs[1].text)
# A plain paragraph having some bold and some italic.

p = d.paragraphs[1].text

# In the paragraph above, we have these run objects, 
# run objects are defined each time the style changes

# First run is 'A plain paragraph having some'
# Second run since it changes to bold: 'bold'
# Third run since it changes back to regular: 'and some'
# Fourth run for italic change: 'italic

p = d.paragraphs[1]
print(p.runs[0].text)
# A plain paragraph having some
print(p.runs[1].text)
# bold
print(p.runs[2].text)
#  and some
print(p.runs[3].text)
# italic

# You can also check if a run is bolded, italic:
print(p.runs[3].italic)
# True
print(p.runs[0].italic)
# None 

p.runs[3].underline = True
p.runs[3].text = 'ITALIC and underlined.'

p.style = 'Title'

d.save('.\\demo2.docx')

