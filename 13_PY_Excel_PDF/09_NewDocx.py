# Creating brand new docx files 

import docx

# To create a brand new doc use the Document()
# method
d = docx.Document()


d.add_paragraph("Hello from Python! This paragraph\
                was created with .add_paragraph method")

d.add_paragraph('Another paragraph')

p = d.paragraphs[0]
p.add_run('This is a new run')
p.runs[1].bold = True

d.save('.\\brand_new_docx_file.docx')